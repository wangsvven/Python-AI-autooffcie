#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
脚本功能：Word文档批量填充Excel数据（统一格式：宋体10号）
核心特性：
1. 支持单/多Word模板批量处理，按"桩号"生成独立Word文件
2. 双填充模式：通用占位符填充 + 表格坐标填充，均强制宋体10号
3. 智能格式处理：日期自动格式化、数值优化、空值统一显示"/"
4. 结构化代码：清晰划分配置区、功能区等模块，易维护
作者：编程助手
更新时间：2026-02-05
适配场景：资料填充
"""

# ==============================================================================
# 【1. 核心库导入区】- 仅导入必要库，注释说明用途
# ==============================================================================
import pandas as pd  # 数据处理：读取Excel、数据格式化
from docx import Document  # Word操作：读写文档、操作段落/表格
from docx.shared import Pt  # Word格式：字体大小设置
from docx.oxml.ns import qn  # Word格式：中文字体兼容（解决宋体显示问题）
from docx.enum.text import WD_ALIGN_PARAGRAPH  # Word格式：文本对齐方式
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT  # Word格式：单元格垂直对齐
import os  # 系统操作：路径处理、文件夹创建
from datetime import datetime  # 日期处理：日期解析与格式化
import re  # 文本处理：正则匹配、日期提取


# ==============================================================================
# 【2. 核心配置区】- 仅需修改此区域，其余代码无需改动
# 注释示例说明：
# 1. 占位符填充示例：{"{{桩号}}": "桩号"} 表示将Word中{{桩号}}替换为Excel"桩号"列的值
# 2. 表格坐标填充示例：{"实测偏差": (5, 25)} 表示将Excel"实测偏差"列填入Word第一个表格第6行第26列
# 3. 日期格式化示例：{"施工日期": "%Y年%m月%d日"} 表示将Excel"施工日期"列格式化为"2024年05月20日"
# 4. 单位配置示例：{"呼称高": "m"} 表示将数值5处理为"5m"
# ==============================================================================
class Config:
    """配置类：集中管理所有可配置项，便于维护和修改"""

    # -------------------------- 路径配置 --------------------------
    # Excel数据文件路径（绝对路径/相对路径）
    EXCEL_FILE = ''
    SHEET_NAME = 'Sheet2'  # Excel工作表名称
    WORD_TEMPLATE = ''  # 单模板路径
    WORD_TEMPLATE_FOLDER = ''  # 多模板文件夹（优先级高于单模板）
    OUTPUT_FOLDER = './填充结果/'  # 输出文件夹（自动创建）

    # -------------------------- 业务配置 --------------------------
    PRIMARY_KEY = '桩号'  # 数据匹配主键（按此列生成文件）
    OUTPUT_FILE_SUFFIX = ''  # 输出文件后缀（如"_填充完成"，最终文件名为"桩号_填充完成.docx"）

    # -------------------------- 填充规则配置 --------------------------
    # 1. 表格坐标填充：{Excel列名: (表格行索引, 表格列索引)}（索引从0开始）
    # 示例：'实测偏差': (5, 25) → 第一个表格第6行第26列填充Excel"实测偏差"列数据
    TABLE_CELL_MAP = {
        '塔型': (0, 7)    # Excel"设计桩号"列 → Word表格第0行第7列
        # 更多示例：
        # '基础埋深': (6, 10),    # Excel"基础埋深"列 → Word表格第6行第10列
    }

    # 2. 通用占位符填充：{Word占位符: Excel列名}（支持段落/表格任意位置）
    # 示例："{{施工单位}}": "施工单位名称" → 将Word中{{施工单位}}替换为Excel"施工单位名称"列的值
    PLACEHOLDER_MAP = {
        # "{{桩号}}": "桩号",  # Word{{桩号}} → Excel"桩号"列
        # "{{施工日期}}": "施工日期"  # Word{{施工日期}} → Excel"施工日期"列
    }

    # -------------------------- 格式配置 --------------------------
    # 全局字体设置（所有填充内容强制此格式）
    FONT_NAME = '宋体'  # 字体名称（如"微软雅黑"、"黑体"）
    FONT_SIZE = Pt(10)  # 字体大小（10号）
    CELL_ALIGNMENT = WD_ALIGN_PARAGRAPH.CENTER  # 表格单元格对齐方式（居中）

    # 日期格式化：{需格式化的列名: 输出格式}
    # 示例：'施工日期': '%Y年%m月%d日' → 2024-05-20 → 2024年05月20日
    DATE_FORMAT_MAP = {
        '施工日期': '%Y年%m月%d日',
        '检查日期': '%Y年%m月%d日'
        # 更多示例：
        # '验收日期': '%Y/%m/%d',  # 2024/05/20
        # '开工日期': '%m月%d日',   # 05月20日
    }

    # 数值单位配置：{Excel列名: 单位}（如5 → 5m）
    # 示例：'呼称高': 'm' → 5 → 5m；'紧线后': '%' → 0.5 → 0.5%
    UNIT_MAP = {
        '呼称高': 'm',
        '塔全高': 'm',
        '放线前': 'm',
        '紧线后': '%',
        '直线塔结构倾斜': '%'
    }

    # 数值优化列：自动去除末尾0（如5.0→5，5.10→5.1）
    OPTIMIZE_DECIMAL_COLUMNS = ['呼称高', '塔全高']


# ==============================================================================
# 【3. 工具函数区】- 独立封装通用功能，便于复用和调试
# ==============================================================================
class ExcelDataProcessor:
    """Excel数据处理工具类"""

    @staticmethod
    def format_date(value, target_format):
        """
        日期格式化：支持多种输入格式，空值返回"/"
        :param value: 原始日期值（datetime/字符串/数字）
        :param target_format: 目标格式（如%Y年%m月%d日）
        :return: 格式化后的日期字符串
        """
        # 空值处理
        if pd.isna(value) or value == '' or str(value).strip() == 'nan':
            return "/"

        # datetime类型直接格式化
        if isinstance(value, (pd.Timestamp, datetime)):
            return value.strftime(target_format)

        # 清理字符串
        val_str = str(value).strip().split(' ')[0]  # 去掉时间部分

        # 处理Excel日期序列号
        if val_str.replace('.', '').isdigit():
            try:
                days = float(val_str)
                base_date = datetime(1899, 12, 30) if days >= 60 else datetime(1899, 12, 30)
                date_obj = base_date + pd.Timedelta(days=days)
                return date_obj.strftime(target_format)
            except Exception as e:
                print(f"⚠️ 日期转换警告：{val_str} → {str(e)[:50]}")
                return val_str

        # 尝试解析常见日期格式
        date_patterns = ['%Y-%m-%d', '%Y/%m/%d', '%Y年%m月%d日', '%m/%d/%Y', '%d/%m/%Y']
        for pattern in date_patterns:
            try:
                return datetime.strptime(val_str, pattern).strftime(target_format)
            except:
                continue

        # 从文本中提取年月日
        try:
            year = re.findall(r'(\d{4})年', val_str)
            month = re.findall(r'(\d{1,2})月', val_str)
            day = re.findall(r'(\d{1,2})日', val_str)
            if year and month and day:
                return f"{year[0]}年{month[0].zfill(2)}月{day[0].zfill(2)}日"
        except:
            pass

        print(f"⚠️ 日期解析失败：{val_str}（返回原值）")
        return val_str

    @staticmethod
    def optimize_number(value):
        """
        数值优化：去除末尾多余0，空值返回"/"
        :param value: 原始数值（float/int/字符串）
        :return: 优化后的字符串
        """
        if pd.isna(value):
            return "/"

        try:
            num = float(value)
            if num.is_integer():
                return str(int(num))
            return str(num).rstrip('0').rstrip('.') if '.' in str(num) else str(num)
        except:
            return str(value)

    @staticmethod
    def load_excel_data(config):
        """
        加载并验证Excel数据
        :param config: 配置类实例
        :return: 清理后的DataFrame
        """
        # 检查Excel文件是否存在
        if not os.path.exists(config.EXCEL_FILE):
            raise FileNotFoundError(f"Excel文件不存在：{config.EXCEL_FILE}")

        # 读取Excel
        df = pd.read_excel(config.EXCEL_FILE, sheet_name=config.SHEET_NAME)
        df.columns = df.columns.str.strip()  # 清理列名空格

        # 验证主键列
        if config.PRIMARY_KEY not in df.columns:
            raise ValueError(f"Excel缺少主键列：{config.PRIMARY_KEY} | 现有列：{list(df.columns)}")

        # 验证占位符对应列
        missing_cols = [col for _, col in config.PLACEHOLDER_MAP.items() if col not in df.columns]
        if missing_cols:
            raise ValueError(f"Excel缺少列：{missing_cols} | 现有列：{list(df.columns)}")

        print(f"✅ 成功读取Excel：{len(df)}行数据，{len(df.columns)}列字段")
        return df


class WordFormatter:
    """Word格式处理工具类"""

    @staticmethod
    def set_font_style(run, config):
        """
        统一设置字体样式（核心：强制宋体10号）
        :param run: Word的Run对象
        :param config: 配置类实例
        """
        run.font.name = config.FONT_NAME  # 设置英文字体
        run.font.size = config.FONT_SIZE  # 设置字体大小
        run._element.rPr.rFonts.set(qn('w:eastAsia'), config.FONT_NAME)  # 中文字体兼容

    @staticmethod
    def fill_table_cell(cell, text, config):
        """
        填充表格单元格并设置格式
        :param cell: Word单元格对象
        :param text: 填充文本
        :param config: 配置类实例
        """
        # 清空单元格
        cell.text = ""

        # 设置单元格格式
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        para.alignment = config.CELL_ALIGNMENT

        # 添加文本并设置字体
        run = para.add_run(str(text))
        WordFormatter.set_font_style(run, config)

    @staticmethod
    def replace_placeholders(doc, data, config):
        """
        替换所有占位符并强制设置宋体10号
        :param doc: Word文档对象
        :param data: 单行数据字典
        :param config: 配置类实例
        """
        # 收集所有段落（表格内+表格外）
        all_paragraphs = []
        all_paragraphs.extend(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_paragraphs.extend(cell.paragraphs)

        # 遍历替换每个占位符
        for placeholder, excel_col in config.PLACEHOLDER_MAP.items():
            # 获取并格式化值
            raw_value = data.get(excel_col, "")
            if pd.isna(raw_value):
                replace_text = "/"
            elif excel_col in config.DATE_FORMAT_MAP:
                replace_text = ExcelDataProcessor.format_date(raw_value, config.DATE_FORMAT_MAP[excel_col])
            elif excel_col in config.OPTIMIZE_DECIMAL_COLUMNS:
                replace_text = ExcelDataProcessor.optimize_number(raw_value)
                if excel_col in config.UNIT_MAP:
                    replace_text += config.UNIT_MAP[excel_col]
            else:
                replace_text = str(raw_value) if not pd.isna(raw_value) else "/"

            # 逐Run替换（保留原有格式，仅修改字体）
            for para in all_paragraphs:
                run_processed = False
                for run in para.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, replace_text)
                        WordFormatter.set_font_style(run, config)
                        run_processed = True
                        break

                # 兜底：段落整体替换
                if not run_processed and placeholder in para.text:
                    para.text = para.text.replace(placeholder, replace_text)
                    for run in para.runs:
                        WordFormatter.set_font_style(run, config)


# ==============================================================================
# 【4. 核心执行区】- 主逻辑入口，调用工具函数完成填充
# ==============================================================================
class WordFiller:
    """Word填充主类"""

    def __init__(self, config):
        self.config = config
        self._prepare_output_folder()

    def _prepare_output_folder(self):
        """创建输出文件夹"""
        if not os.path.exists(self.config.OUTPUT_FOLDER):
            os.makedirs(self.config.OUTPUT_FOLDER)
            print(f"✅ 创建输出文件夹：{self.config.OUTPUT_FOLDER}")

    def _get_word_templates(self):
        """获取Word模板路径列表"""
        # 优先使用多模板文件夹
        if self.config.WORD_TEMPLATE_FOLDER and os.path.exists(self.config.WORD_TEMPLATE_FOLDER):
            templates = [
                os.path.join(self.config.WORD_TEMPLATE_FOLDER, f)
                for f in os.listdir(self.config.WORD_TEMPLATE_FOLDER)
                if f.endswith('.docx') and not f.startswith('~$')
            ]
            if templates:
                print(f"✅ 加载多模板：共{len(templates)}个文件")
                return templates

        # 使用单模板
        if self.config.WORD_TEMPLATE and os.path.exists(self.config.WORD_TEMPLATE):
            print(f"✅ 加载单模板：{self.config.WORD_TEMPLATE}")
            return [self.config.WORD_TEMPLATE]

        raise FileNotFoundError("未找到有效Word模板文件")

    def _format_cell_value(self, excel_col, raw_val, config):
        """
        统一格式化单元格值（抽离重复逻辑，提升可维护性）
        :param excel_col: Excel列名
        :param raw_val: 原始值
        :param config: 配置实例
        :return: 格式化后的文本
        """
        if pd.isna(raw_val):
            return "/"
        elif excel_col in config.DATE_FORMAT_MAP:
            return ExcelDataProcessor.format_date(raw_val, config.DATE_FORMAT_MAP[excel_col])
        elif excel_col in config.OPTIMIZE_DECIMAL_COLUMNS:
            val = ExcelDataProcessor.optimize_number(raw_val)
            return val + config.UNIT_MAP.get(excel_col, "")
        else:
            return str(raw_val)

    def process_single_station(self, template_path, station, data_row):
        """
        处理单个桩号的数据填充
        :param template_path: Word模板路径
        :param station: 桩号名称
        :param data_row: 单行数据字典
        """
        # 构建输出路径
        station_clean = str(station).strip()
        output_path = os.path.join(
            self.config.OUTPUT_FOLDER,
            f"{station_clean}{self.config.OUTPUT_FILE_SUFFIX}.docx"
        )

        try:
            # 打开模板
            doc = Document(template_path)

            # 步骤1：替换占位符（强制宋体10号）
            WordFormatter.replace_placeholders(doc, data_row, self.config)

            # 步骤2：填充表格坐标（强制宋体10号）
            if doc.tables:
                main_table = doc.tables[0]  # 取第一个表格
                for excel_col, (row_idx, col_idx) in self.config.TABLE_CELL_MAP.items():
                    # 跳过不存在的列
                    if excel_col not in data_row:
                        print(f"⏩ 跳过[{station_clean}]：缺少列{excel_col}")
                        continue

                    # 检查行列索引是否越界
                    if row_idx >= len(main_table.rows) or col_idx >= len(main_table.columns):
                        print(f"⏩ 跳过[{station_clean}]：表格行列越界（行{row_idx + 1}，列{col_idx + 1}）")
                        continue

                    # 格式化值
                    fill_text = self._format_cell_value(excel_col, data_row[excel_col], self.config)

                    # 填充单元格
                    WordFormatter.fill_table_cell(main_table.cell(row_idx, col_idx), fill_text, self.config)

            # 保存文件
            doc.save(output_path)
            print(f"✅ 成功[{station_clean}]：{os.path.basename(output_path)}")

        except Exception as e:
            print(f"❌ 失败[{station_clean}]：{str(e)[:80]}")

    def run(self):
        """主执行函数"""
        try:
            # 1. 加载Excel数据
            df = ExcelDataProcessor.load_excel_data(self.config)

            # 2. 获取Word模板
            templates = self._get_word_templates()

            # 3. 遍历每个模板
            for template in templates:
                template_name = os.path.basename(template)
                print(f"\n========== 处理模板：{template_name} ==========")

                # 4. 遍历每个桩号
                unique_stations = df[self.config.PRIMARY_KEY].unique()
                for station in unique_stations:
                    if pd.isna(station) or str(station).strip() == "":
                        print(f"⏩ 跳过：空桩号")
                        continue

                    # 获取当前桩号数据
                    station_data = df[df[self.config.PRIMARY_KEY] == station].iloc[0].to_dict()
                    # 处理单个桩号
                    self.process_single_station(template, station, station_data)

            # 完成提示
            print(f"\n🎉 全部处理完成！")
            print(f"📁 输出目录：{os.path.abspath(self.config.OUTPUT_FOLDER)}")
            print(f"📌 格式说明：所有填充内容均为{self.config.FONT_NAME}{self.config.FONT_SIZE.pt}号字体")

        except Exception as e:
            print(f"\n❌ 执行失败：{str(e)}")
            raise


# ==============================================================================
# 【5. 脚本入口】- 一键执行
# ==============================================================================
if __name__ == "__main__":
    # 创建配置实例
    config = Config()
    # 创建填充实例并执行
    filler = WordFiller(config)
    filler.run()
