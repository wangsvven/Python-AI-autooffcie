###  BridgePile_AutoFill_Pro.py

import pandas as pd  # 用于处理 Excel 数据的工具
from docx import Document  # 用于处理 Word 文档的工具
from docx.shared import Pt  # 用于设置字号大小
from docx.oxml.ns import qn  # 用于设置中文字体兼容性
from docx.enum.text import WD_ALIGN_PARAGRAPH  # 用于设置水平居中
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT  # 用于设置垂直居中
import os  # 用于处理文件路径和文件夹

# ============================================================
# 【第一部分：小白配置区】—— 每次换项目，只需改这里的文字
# ============================================================

# 1. 文件夹路径：请把空白 Word 模板放在这个路径下
# 注意：路径最后的名称必须和电脑上的文件夹名一致
INPUT_WORD_FOLDER = '/Users/mac/Library/CloudStorage/OneDrive-个人/1.项目/攀枝花米易撒莲丙谷光伏发电项目（35kV 集电线路）/6.过程资料/7.相关数据/表D.0.4 灌注桩基础检查记录表'

# 2. Excel 数据库：你的数据来源表
EXCEL_DATABASE = '/Users/mac/Desktop/work/串筒灌注记录_最终完美版.xlsx'

# 3. 输出位置：填充好的文件会放在脚本所在文件夹的这个新文件夹里
OUTPUT_FOLDER = './填充结果_新生成/'

# 4. 匹配依据：Excel 里哪一列的文字对应 Word 的文件名？
STATION_COLUMN_NAME = '设计桩号'

# 5. 坐标映射：【Excel表头名】对应【Word表格里的第几列】
# 注意：这里的数字是通过“坐标探测脚本”得到的物理列号
COLUMN_MAP = {
    '灌1': 0,
    '拆2': 3,
    '斗3': 5,
    '折4': 9,
    '孔5': 13,
    '拆6': 16,
    '埋7': 22
}

# 6. 行数设置
START_ROW_INDEX = 16  # 数据从 Word 表格的第 16 行（代码索引16）开始填入
MAX_ROWS_TO_FILL = 15  # 表格数据区共有 15 行（即填到第 30 行截止）

# 7. 文件后缀：填 "" 表示保持原名，填 "_已填充" 会在文件名后加注
FILE_SUFFIX = ""


# ============================================================
# 【第二部分：核心功能区】—— 负责改字体、对齐和填数，建议不要修改
# ============================================================

def fill_cell_with_font_style(cell, text):
    """
    此函数负责：把文字填进去，并强制设为 宋体 10号 居中
    """
    # 让单元格内容上下居中
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # 清空单元格原本的内容
    cell.text = ""

    # 获取单元格里的段落，如果没有就新建一个
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()

    # 让文字左右居中
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 创建文字块并填入内容
    run = para.add_run(str(text))

    # 设置字号为 10 磅
    run.font.size = Pt(10)

    # 强制设置中文字体为“宋体” (兼容 Word 核心协议)
    run.font.name = '宋体'  # 这一行设置西文字体
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 这一行锁定中文字体


def run_universal_filler():
    """
    主程序：负责批量读写文件
    """
    # 如果没有输出文件夹，就自动新建一个
    if not os.path.exists(OUTPUT_FOLDER):
        os.makedirs(OUTPUT_FOLDER)

    # 检查 Excel 文件在不在
    if not os.path.exists(EXCEL_DATABASE):
        print(f"【错误】找不到 Excel 文件，请检查路径是否正确: {EXCEL_DATABASE}")
        return

    # 读取 Excel 内容
    df = pd.read_excel(EXCEL_DATABASE)
    # 自动把表头前后的空格删掉，防止匹配出错
    df.columns = df.columns.str.strip()

    # 检查 Excel 里有没有“设计桩号”这一列
    if STATION_COLUMN_NAME not in df.columns:
        print(f"【错误】Excel 里没有找到 '{STATION_COLUMN_NAME}' 这一列")
        return

    # 找到 Excel 里所有不重复的桩号，一个个开始处理
    all_stations = df[STATION_COLUMN_NAME].unique()
    print(f"--- 发现 {len(all_stations)} 个桩号，开始批量生产... ---")

    for station in all_stations:
        # 确定 Word 模板的文件名
        original_file_name = f"{str(station).strip()}.docx"
        new_file_name = f"{str(station).strip()}{FILE_SUFFIX}.docx"

        # 拼接完整路径
        input_path = os.path.join(INPUT_WORD_FOLDER, original_file_name)
        output_path = os.path.join(OUTPUT_FOLDER, new_file_name)

        # 如果找不到对应的 Word 模板，就跳过
        if not os.path.exists(input_path):
            print(f"【跳过】文件夹里没找到模板: {original_file_name}")
            continue

        try:
            # 打开 Word 文档
            doc = Document(input_path)
            # 找到文档里的第一个表格
            table = doc.tables[0]

            # 从 Excel 里筛选出属于这个桩号的所有行
            station_data = df[df[STATION_COLUMN_NAME] == station].reset_index(drop=True)

            # 开始填数（按行遍历）
            for i in range(len(station_data)):
                # 如果 Excel 数据多于 15 行，就只填前 15 行，防止撑破表格
                if i >= MAX_ROWS_TO_FILL:
                    break

                excel_row = station_data.iloc[i]

                # 开始填每一列的数据
                for excel_col, word_idx in COLUMN_MAP.items():
                    if excel_col in df.columns:
                        # 获取 Excel 里的数值
                        val = excel_row[excel_col]
                        # 如果是空的，就填个斜杠 "/"；否则转成文字
                        content = str(val) if pd.notna(val) else "/"

                        # 找到对应的 Word 单元格并填入
                        target_cell = table.cell(START_ROW_INDEX + i, word_idx)
                        fill_cell_with_font_style(target_cell, content)

            # 全部填完，保存到新文件夹里
            doc.save(output_path)
            print(f"【成功】已生成: {new_file_name}")

        except Exception as e:
            # 如果中间出错了（比如 Word 被占用），报错并继续下一个
            print(f"【异常】处理 {station} 时出错: {e}")

    print(f"\n恭喜！所有文件已完成，请去这里查看：{OUTPUT_FOLDER}")


# --- 脚本入口 ---
if __name__ == "__main__":
    run_universal_filler()
