#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
================================================================================
【脚本名称】：Word文档批量填充Excel数据 (v3.5 终极注释版)
【应用场景】：工程管理资料自动化（如：攀枝花米易撒莲丙谷光伏发电项目35kV集电线路资料）
【核心能力】：
    1. 自动读取 Excel 表格，按“设计桩号”批量生成 Word 文档。
    2. 支持 3 种填充模式：表格固定坐标、{占位符} 替换、指定关键字后缀（零侵入）。
    3. 智能排版：所有生成的内容强制锁定为“宋体 10号”，自动居中。
    4. 智能数据：自动格式化日期，自动去除数值末尾多余的 0，自动追加物理单位。
    5. 精准控制：支持只打印指定的“桩号名单”，或只打印指定的“Excel行号区间”。
================================================================================
"""

# ==============================================================================
# 【1. 核心库导入区】 - 脚本运行所需的基础工具箱
# ==============================================================================
import pandas as pd  # 数据处理大神：负责读取和切片 Excel 数据
from docx import Document  # Word 操作手：负责打开、修改和保存 Word 文档
from docx.shared import Pt  # 格式工具：负责设置字体大小（Point）
from docx.oxml.ns import qn  # 格式工具：负责解决中文字体（如宋体）在 Word 中的兼容性问题
from docx.enum.text import WD_ALIGN_PARAGRAPH  # 格式工具：负责段落对齐（居中、靠左等）
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT  # 格式工具：负责表格单元格的垂直对齐
import os  # 系统管家：负责创建文件夹、检查文件是否存在
from datetime import datetime  # 时间管理：负责识别和转换各种日期格式
import re  # 文本侦探：正则表达式库，负责从杂乱的文字中提取目标内容（如日期）


# ==============================================================================
# 【2. 核心配置区】 ★★★ 日常使用只需修改这里，下面代码不用动 ★★★
# ==============================================================================
class Config:
    """集中管理所有配置参数，方便日常使用和维护"""

    # -------------------------- A. 文件路径配置 --------------------------
    # 【必填】你的 Excel 数据源文件路径
    EXCEL_FILE = '/Users/mac/Library/CloudStorage/OneDrive-个人/1.项目/攀枝花米易撒莲丙谷光伏发电项目（35kV 集电线路）/6.过程资料/7.相关数据/（数据）表D.0.8 铁塔组立检查记录表.xlsx'

    # 【必填】Excel 里面具体要读取的工作表名称（如 Sheet1, Sheet2）
    SHEET_NAME = 'Sheet2'

    # 【必填】你的 Word 模板文件路径
    WORD_TEMPLATE = '/Users/mac/Desktop/work/表D.0.8 铁塔组立检查记录表py.docx'

    # 【选填】如果你有一堆模板都在一个文件夹里，可以填这里，留空则只用上面的单模板
    WORD_TEMPLATE_FOLDER = ''

    # 【必填】生成的新 Word 文档保存在哪里？（文件夹不存在会自动创建）
    OUTPUT_FOLDER = './填充结果a/'

    # -------------------------- B. 业务基础配置 --------------------------
    # 【必填】数据的主心骨（通常是第一列，脚本会根据这一列的名字来给 Word 文件命名）
    PRIMARY_KEY = '设计桩号'

    # 【选填】生成文件的后缀名（例如填入 '_已完成'，生成的文件名就是 '线塔1_已完成.docx'）
    OUTPUT_FILE_SUFFIX = ''

    # -------------------------- C. ★ 高级生成范围控制（类似打印机设置） --------------------------
    # 模式一：按“具体名称”精确指定。
    # 用法：填入需要生成的桩号，如 ['15号塔', '18号塔']。填 [] 代表全部生成。
    TARGET_STATIONS = []

    # 模式二：按“Excel真实行号区间”指定。
    # 用法：[起始行号, 结束行号]。例如 [3, 10] 表示只生成 Excel 左侧显示的第 3 行到第 10 行。填 [] 代表全部生成。
    TARGET_ROW_RANGE = []

    # -------------------------- D. 填充规则配置 --------------------------
    # 规则 1：【表格坐标填充】
    # 格式：'Excel表头名': (Word表格的行号, Word表格的列号) —— 注意：行号列号从 0 开始算！
    TABLE_CELL_MAP = {
        '设计桩号': (1, 3),  # 把"设计桩号"填入第 2 行第 4 列
        '杆塔型': (1, 8),  # 把"杆塔型"填入第 2 行第 9 列
        '呼称高': (0, 12),
        '塔全高': (1, 12),
        '施工日期': (0, 19),
        '检查日期': (1, 19),
        '直线塔结构倾斜': (16, 19),
        '放线前': (21, 19),
        '紧线后': (22, 19)
    }

    # 规则 2：【占位符填充】（需在 Word 模板里提前写好，如 {项目名称}）
    # 格式：'{Word里的占位符}': 'Excel表头名'
    PLACEHOLDER_MAP = {
    }

    # 规则 3：【关键字追加填充】（零侵入式，不用改模板）
    # 格式：'Word里固定的文字': 'Excel表头名'
    KEYWORD_APPEND_MAP = {
        '编号：': '编号'  # 脚本会在Word里找到"编号："，然后紧贴着后面填入Excel里"编号"列的数据
    }

    # -------------------------- E. 格式化控制中心 --------------------------
    # 全局字体设置：所有程序填进去的字，统统变成这个样式
    FONT_NAME = '宋体'
    FONT_SIZE = Pt(10)  # 10号字体
    CELL_ALIGNMENT = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐

    # 日期变身：自动把各种乱七八糟的日期变成统一格式
    DATE_FORMAT_MAP = {
        '施工日期': '%Y年%m月%d日',  # 最终效果：2026年03月04日
        '检查日期': '%Y年%m月%d日'
    }

    # 单位外衣：给填进去的数字自动穿上单位（只要在这里配置了，哪怕不是去0的列也会加单位）
    UNIT_MAP = {
        '呼称高': 'm',
        '塔全高': 'm',
        '放线前': '%',
        '紧线后': '%',
        '直线塔结构倾斜': '‰'
    }

    # 去零瘦身：自动把 5.0 变成 5，把 5.10 变成 5.1   ['呼称高', '塔全高']
    OPTIMIZE_DECIMAL_COLUMNS = []


# ==============================================================================
# 【3. 工具函数区】 - 脚本的内部发动机，处理各种脏活累活
# ==============================================================================
class ExcelDataProcessor:
    """处理 Excel 数据的各种疑难杂症"""

    @staticmethod
    def format_date(value, target_format):
        """核心逻辑：把 Excel 里千奇百怪的日期（时间戳、字符串、数字）统一成标准格式"""
        # 如果是空单元格，统一填个斜杠
        if pd.isna(value) or value == '' or str(value).strip() == 'nan':
            return "/"

        # 已经是标准时间对象的，直接按格式输出
        if isinstance(value, (pd.Timestamp, datetime)):
            return value.strftime(target_format)

        val_str = str(value).strip().split(' ')[0]  # 扔掉具体时分秒，只留年月日

        # 应对 Excel 的“数字日期”特性（Excel有时会把日期存为距1899年的天数）
        if val_str.replace('.', '').isdigit():
            try:
                days = float(val_str)
                base_date = datetime(1899, 12, 30) if days >= 60 else datetime(1899, 12, 30)
                date_obj = base_date + pd.Timedelta(days=days)
                return date_obj.strftime(target_format)
            except Exception as e:
                print(f"⚠️ 日期转换警告：{val_str} → {str(e)[:50]}")
                return val_str

        # 暴力尝试各种常见的日期格式组合
        date_patterns = ['%Y-%m-%d', '%Y/%m/%d', '%Y年%m月%d日', '%m/%d/%Y', '%d/%m/%Y']
        for pattern in date_patterns:
            try:
                return datetime.strptime(val_str, pattern).strftime(target_format)
            except:
                continue

        # 终极大法：如果上面都不行，用正则把数字抠出来硬拼
        try:
            year = re.findall(r'(\d{4})年', val_str)
            month = re.findall(r'(\d{1,2})月', val_str)
            day = re.findall(r'(\d{1,2})日', val_str)
            if year and month and day:
                return f"{year[0]}年{month[0].zfill(2)}月{day[0].zfill(2)}日"
        except:
            pass

        return val_str  # 实在解析不出来，就把原样文字扔进去

    @staticmethod
    def optimize_number(value):
        """核心逻辑：消灭无意义的零（比如让 5.0 变 5）"""
        if pd.isna(value):
            return "/"
        try:
            num = float(value)
            # 如果是整数，直接转成 int 形式（去掉.0）
            if num.is_integer():
                return str(int(num))
            # 否则去掉末尾的 0 和多余的小数点
            return str(num).rstrip('0').rstrip('.') if '.' in str(num) else str(num)
        except:
            return str(value)  # 如果原本就是中文（如"不适用"），原样返回

    @staticmethod
    def load_excel_data(config):
        """读取数据并做基本体检"""
        if not os.path.exists(config.EXCEL_FILE):
            raise FileNotFoundError(f"救命，Excel文件没找到：{config.EXCEL_FILE}")

        df = pd.read_excel(config.EXCEL_FILE, sheet_name=config.SHEET_NAME)
        df.columns = df.columns.str.strip()  # 去掉表头里不小心敲进去的空格

        # 检查必须存在的列，防止运行一半崩溃
        if config.PRIMARY_KEY not in df.columns:
            raise ValueError(f"Excel里找不到主键列[{config.PRIMARY_KEY}]")

        print(f"✅ 成功读取Excel：包含 {len(df)} 条有效数据")
        return df


class WordFormatter:
    """Word 文档的美容师：负责往里面填字，并控制长相"""

    @staticmethod
    def set_font_style(run, config):
        """底层逻辑：给被选中的字强制套上宋体 10 号"""
        run.font.name = config.FONT_NAME
        run.font.size = config.FONT_SIZE
        # 中文字体必须这么设置才能在 Word 里生效
        run._element.rPr.rFonts.set(qn('w:eastAsia'), config.FONT_NAME)

    @staticmethod
    def fill_table_cell(cell, text, config):
        """底层逻辑：往表格具体的格子里填字"""
        cell.text = ""  # 先把格子里原来的东西清空
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 上下居中
        para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        para.alignment = config.CELL_ALIGNMENT  # 左右居中
        run = para.add_run(str(text))
        WordFormatter.set_font_style(run, config)

    @staticmethod
    def replace_placeholders(doc, data, config):
        """底层逻辑：全篇扫描 {占位符} 并替换"""
        # 把文档里所有的段落（表格里的、表格外的）全部收集起来
        all_paragraphs = []
        all_paragraphs.extend(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_paragraphs.extend(cell.paragraphs)

        for placeholder, excel_col in config.PLACEHOLDER_MAP.items():
            raw_value = data.get(excel_col, "")

            # 拿到数据后，先根据配置进行格式化、去0、加单位
            if pd.isna(raw_value):
                replace_text = "/"
            elif excel_col in config.DATE_FORMAT_MAP:
                replace_text = ExcelDataProcessor.format_date(raw_value, config.DATE_FORMAT_MAP[excel_col])
            else:
                if excel_col in config.OPTIMIZE_DECIMAL_COLUMNS:
                    replace_text = ExcelDataProcessor.optimize_number(raw_value)
                else:
                    replace_text = str(raw_value) if not pd.isna(raw_value) else "/"

                # 只要配了单位，就得加（跳过空值）
                if excel_col in config.UNIT_MAP and replace_text != "/":
                    replace_text += config.UNIT_MAP[excel_col]

            # 开始全篇搜索替换
            for para in all_paragraphs:
                run_processed = False
                for run in para.runs:  # 尽量在最小单元(Run)替换，以保留原有格式
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, replace_text)
                        WordFormatter.set_font_style(run, config)
                        run_processed = True
                        break
                # 如果被 Word 底层强行切断了，就整个段落暴力替换
                if not run_processed and placeholder in para.text:
                    para.text = para.text.replace(placeholder, replace_text)
                    for run in para.runs:
                        WordFormatter.set_font_style(run, config)

    @staticmethod
    def append_keywords(doc, data, config, format_func):
        """底层逻辑：找关键字（如“编号：”），然后在它屁股后面追加数据"""
        all_paragraphs = []
        all_paragraphs.extend(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_paragraphs.extend(cell.paragraphs)

        for keyword, excel_col in config.KEYWORD_APPEND_MAP.items():
            if excel_col not in data:
                continue

            raw_value = data.get(excel_col, "")
            # 调用主类的格式化大师，处理去0加单位等事务
            append_text = format_func(excel_col, raw_value, config)

            # 要替换成的最终效果 = "编号：" + "X塔数据"
            target_replace = keyword + append_text

            for para in all_paragraphs:
                # 如果段落里有"编号："，并且还没被追加过数据，就开干
                if keyword in para.text and target_replace not in para.text:
                    run_processed = False
                    for run in para.runs:
                        if keyword in run.text:
                            run.text = run.text.replace(keyword, target_replace)
                            WordFormatter.set_font_style(run, config)
                            run_processed = True
                            break
                    if not run_processed:
                        para.text = para.text.replace(keyword, target_replace)
                        for run in para.runs:
                            WordFormatter.set_font_style(run, config)


# ==============================================================================
# 【4. 核心执行区】 - 脚本的大脑指挥中心，统筹全局
# ==============================================================================
class WordFiller:
    def __init__(self, config):
        self.config = config
        self._prepare_output_folder()

    def _prepare_output_folder(self):
        """确保输出文件夹乖乖躺在那里"""
        if not os.path.exists(self.config.OUTPUT_FOLDER):
            os.makedirs(self.config.OUTPUT_FOLDER)
            print(f"✅ 创建输出文件夹：{self.config.OUTPUT_FOLDER}")

    def _get_word_templates(self):
        """寻找 Word 模板，支持单文件或整个文件夹"""
        if self.config.WORD_TEMPLATE_FOLDER and os.path.exists(self.config.WORD_TEMPLATE_FOLDER):
            templates = [
                os.path.join(self.config.WORD_TEMPLATE_FOLDER, f)
                for f in os.listdir(self.config.WORD_TEMPLATE_FOLDER)
                if f.endswith('.docx') and not f.startswith('~$')  # 避开正在被打开的隐藏缓存文件
            ]
            if templates:
                print(f"✅ 加载多模板：共 {len(templates)} 个文件")
                return templates
        if self.config.WORD_TEMPLATE and os.path.exists(self.config.WORD_TEMPLATE):
            print(f"✅ 加载单模板：{self.config.WORD_TEMPLATE}")
            return [self.config.WORD_TEMPLATE]
        raise FileNotFoundError("未找到有效Word模板文件")

    def _format_cell_value(self, excel_col, raw_val, config):
        """统一对数值进行格式化处理（给上面三种填充模式共用）"""
        if pd.isna(raw_val):
            return "/"
        elif excel_col in config.DATE_FORMAT_MAP:
            return ExcelDataProcessor.format_date(raw_val, config.DATE_FORMAT_MAP[excel_col])
        else:
            if excel_col in config.OPTIMIZE_DECIMAL_COLUMNS:
                val = ExcelDataProcessor.optimize_number(raw_val)
            else:
                val = str(raw_val)
            if excel_col in config.UNIT_MAP and val != "/":
                val += config.UNIT_MAP[excel_col]
            return val

    def process_single_station(self, template_path, station, data_row):
        """生成一根指定“桩号”的文档（核心组装流水线）"""
        station_clean = str(station).strip()
        # 组装最终存盘的路径名：输出目录 / 桩号名字 + 后缀 .docx
        output_path = os.path.join(
            self.config.OUTPUT_FOLDER,
            f"{station_clean}{self.config.OUTPUT_FILE_SUFFIX}.docx"
        )

        try:
            doc = Document(template_path)  # 打开模具

            # 工序 1：把里面的 {项目名称} 这种暗号替换掉
            WordFormatter.replace_placeholders(doc, data_row, self.config)

            # 工序 2：找到“编号：”这种暗号，在后面默默补上内容
            WordFormatter.append_keywords(doc, data_row, self.config, self._format_cell_value)

            # 工序 3：定位到表格第 X 行第 Y 列，精准打入数据
            if doc.tables:
                main_table = doc.tables[0]  # 默认操作文档里的第一个表格
                for excel_col, (row_idx, col_idx) in self.config.TABLE_CELL_MAP.items():
                    if excel_col not in data_row:
                        continue
                    if row_idx >= len(main_table.rows) or col_idx >= len(main_table.columns):
                        continue
                    fill_text = self._format_cell_value(excel_col, data_row[excel_col], self.config)
                    WordFormatter.fill_table_cell(main_table.cell(row_idx, col_idx), fill_text, self.config)

            doc.save(output_path)  # 生成脱模
            print(f"✅ 成功[{station_clean}]：{os.path.basename(output_path)}")

        except Exception as e:
            print(f"❌ 失败[{station_clean}]：{str(e)[:80]}")

    def run(self):
        """总导演开机：控制整体流程"""
        try:
            # 1. 把 Excel 拖过来
            df = ExcelDataProcessor.load_excel_data(self.config)

            # ---------------- 拦截器 1：按行号精准切片 ----------------
            if self.config.TARGET_ROW_RANGE and len(self.config.TARGET_ROW_RANGE) == 2:
                start_row, end_row = self.config.TARGET_ROW_RANGE
                # 换算逻辑：Excel显示的第 1 行通常是表头，真正的数据从第 2 行开始。
                # 在 Pandas 语言里，数据的第一行索引是 0。
                # 所以要拿真实行号减掉 2，算出计算机能懂的起始索引。
                start_idx = max(0, start_row - 2)
                end_idx = end_row - 1
                df = df.iloc[start_idx:end_idx]
                print(f"🎯 开启【行号打印模式】：已截取 Excel 第 {start_row} 行至第 {end_row} 行的数据")

            templates = self._get_word_templates()

            # ---------------- 拦截器 2：播报名单模式 ----------------
            if self.config.TARGET_STATIONS:
                print(f"🎯 开启【名单打印模式】：仅处理指定名单中的 {len(self.config.TARGET_STATIONS)} 个桩号")

            # 2. 对每个模板，逐行塞入数据
            for template in templates:
                template_name = os.path.basename(template)
                print(f"\n========== 处理模板：{template_name} ==========")

                unique_stations = df[self.config.PRIMARY_KEY].unique()
                for station in unique_stations:
                    if pd.isna(station) or str(station).strip() == "":
                        continue

                    # ---------------- 拦截器 3：按名单过滤 ----------------
                    if self.config.TARGET_STATIONS and station not in self.config.TARGET_STATIONS:
                        continue  # 如果开启了名单模式，且当前人不在这份名单里，直接跳过不干活

                    # 取出属于当前桩号的这一行数据，转成字典方便使用
                    station_data = df[df[self.config.PRIMARY_KEY] == station].iloc[0].to_dict()
                    self.process_single_station(template, station, station_data)

            print(f"\n🎉 全部处理完成！")
            print(f"📁 输出目录：{os.path.abspath(self.config.OUTPUT_FOLDER)}")

        except Exception as e:
            print(f"\n❌ 执行失败：{str(e)}")
            raise


# ==============================================================================
# 【5. 脚本入口】 - 点火器：按运行按钮后，从这里开始点火起飞
# ==============================================================================
if __name__ == "__main__":
    config = Config()  # 把配置单拿到手
    filler = WordFiller(config)  # 把配置单交给执行机器
    filler.run()  # 按下启动按钮
